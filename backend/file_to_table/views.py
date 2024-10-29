from django.shortcuts import render
import os
os.environ['KMP_DUPLICATE_LIB_OK'] = 'True'
import shutil
import fitz
from .models import *
from django.conf import settings
import cv2
from django.shortcuts import render
import layoutparser as lp
import glob
import pandas as pd
import numpy as np
from PIL import Image
from tqdm.auto import tqdm
from blend_modes import divide
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import easyocr
from PIL import Image, ImageFilter
import time
from azure.core.credentials import AzureKeyCredential
from azure.ai.formrecognizer import DocumentAnalysisClient
from dotenv import load_dotenv


load_dotenv()

# Set up Azure Form Recognizer client
key = os.getenv('key')
endpoint = os.getenv('endpoint')

document_analysis_client = DocumentAnalysisClient(endpoint=endpoint, credential=AzureKeyCredential(key))

# Create your views here.

def show_homepage(request):
    return render(request, 'index.html')

def process_image_for_table_extraction(image_path):
    with open(image_path, "rb") as f:
        image_data = f.read()

    poller = document_analysis_client.begin_analyze_document("prebuilt-layout", image_data)
    result = poller.result()

    tables_list = []
    for table in result.tables:
        data = [['' for _ in range(table.column_count)] for _ in range(table.row_count)]
        for cell in table.cells:
            data[cell.row_index][cell.column_index] = cell.content
        df = pd.DataFrame(data)
        df = df.transpose()
        tables_list.append(df)
    return tables_list

def render_table_as_html(tables_list):
    output_html_path = os.path.join(settings.MEDIA_ROOT, "tables.html")
    with open(output_html_path, 'w', encoding='utf-8') as f:
        f.write('<html><head><title>Extracted Tables</title></head><body>\n')
        for i, df in enumerate(tables_list):
            html_table = df.to_html(index=False)
            f.write(f'<h2>Table {i+1}</h2>\n')
            f.write(html_table)
            f.write('<br>\n')
        f.write('</body></html>')
    return output_html_path


def extract_and_display_tables(request):
    image_input = ImageInput.objects.first()  # Assuming you have a model instance with the image
    if not image_input:
        return render(request, 'index.html', {'error': 'No image found.'})
    
    image_path = image_input.image.path
    tables_list = process_image_for_table_extraction(image_path)
    html_path = render_table_as_html(tables_list)
    
    with open(html_path, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    return render(request, 'index.html', {'tables_html': html_content})

###################################################convert to doc starts here#########################################################

from docx import Document
import io
import cv2
from django.http import HttpResponse
from django.core.files.storage import default_storage

def get_y_coordinate(element):
    if hasattr(element, 'boundingRegions') and element.boundingRegions:
        return element.boundingRegions[0].boundingBox[1]
    elif hasattr(element, 'spans') and element.spans:
        return element.spans[0].offset
    return float('inf')

# def add_table_to_doc(table, doc):
#     table_doc = doc.add_table(rows=table.row_count, cols=table.column_count)
#     for cell in table.cells:
#         row = cell.row_index
#         col = cell.column_index
#         table_doc.cell(row, col).text = cell.content

def process_image(image_path):
    image = cv2.imread(image_path)
    doc = Document()

    with open(image_path, "rb") as f:
        image_data = f.read()

    poller = document_analysis_client.begin_analyze_document("prebuilt-layout", image_data)
    result = poller.result()

    page_elements = []
    for page in result.pages:
        if hasattr(page, 'lines'):
            for line in page.lines:
                page_elements.append(('line', get_y_coordinate(line), line))

    for paragraph in result.paragraphs:
        is_table_content = any(cell.content == paragraph.content for table in result.tables for cell in table.cells)
        if not is_table_content:
            page_elements.append(('paragraph', get_y_coordinate(paragraph), paragraph))

    for table in result.tables:
        page_elements.append(('table', get_y_coordinate(table), table))

    # Sort elements by their Y-coordinate to maintain order
    page_elements.sort(key=lambda x: x[1])


    for element_type, y_coord, element in page_elements:
        if element_type == 'paragraph':
            role = element.role
            content = element.content

            if role == "title":
                doc.add_heading(content, level=1)  # Adjust the level as needed
            elif role == "sectionHeading":
                doc.add_heading(content, level=2)  # Adjust the level as needed
            elif role == "pageHeader":
                doc.add_paragraph(content, style='Header')  # Add a header style if desired
            elif role == "pageFooter":
                doc.add_paragraph(content, style='Footer')  # Add a footer style if desired
            elif role == "footnote":
                doc.add_paragraph(content, style='Footnote')  # Add a footnote style if desired
            else:
                doc.add_paragraph(content)  # Default to a regular paragraph

        elif element_type == 'table':
            table_doc = doc.add_table(rows=element.row_count, cols=element.column_count)

            for cell in element.cells:
                row = cell.row_index
                col = cell.column_index
                table_doc.cell(row, col).text = cell.content

    return doc




def convert_image_to_docx(image_path):
    doc = process_image(image_path)

    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

def download_docx(request):
    if request.method == 'POST':
        if 'image' not in request.FILES:
            return HttpResponse("No image file uploaded.", status=400)

        image_file = request.FILES['image']
        image_path = default_storage.save(image_file.name, ContentFile(image_file.read()))
        image_path = os.path.join(settings.MEDIA_ROOT, image_path)

        docx_file = convert_image_to_docx(image_path)
        
        response = HttpResponse(docx_file, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="output.docx"'
        return response
    
    return render(request, 'index.html', {'error': 'Invalid request method.'})
